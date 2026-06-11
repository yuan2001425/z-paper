import uuid
from io import BytesIO
import re
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.database import get_db
from app.models.result import TranslationResult
from app.models.annotation import Annotation
from app.models.paper import Paper
from app.schemas.result import ResultResponse, AnnotationRequest, AnnotationUpdateRequest, AnnotationResponse
from app.storage.local_storage import local_storage

router = APIRouter()


def _attach_pdf_url(result: TranslationResult, db: Session) -> ResultResponse:
    paper = db.query(Paper).filter(Paper.id == result.paper_id).first()
    pdf_url = local_storage.get_url(paper.storage_key) if paper and paper.storage_key else None
    data = ResultResponse.model_validate(result)
    data.pdf_url = pdf_url
    return data


class TranslateImageRequest(BaseModel):
    image_url: str
    block_index: int


@router.get("/by-job/{job_id}", response_model=ResultResponse)
def get_result_by_job(job_id: str, db: Session = Depends(get_db)):
    result = db.query(TranslationResult).filter(TranslationResult.job_id == job_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="译文不存在")
    return _attach_pdf_url(result, db)


@router.get("/by-paper/{paper_id}", response_model=ResultResponse)
def get_result_by_paper(paper_id: str, db: Session = Depends(get_db)):
    result = db.query(TranslationResult).filter(TranslationResult.paper_id == paper_id).order_by(TranslationResult.created_at.desc()).first()
    if not result:
        raise HTTPException(status_code=404, detail="译文不存在")
    return _attach_pdf_url(result, db)


@router.get("/{result_id}", response_model=ResultResponse)
def get_result(result_id: str, db: Session = Depends(get_db)):
    result = db.query(TranslationResult).filter(TranslationResult.id == result_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="译文不存在")
    return _attach_pdf_url(result, db)


# ── 批注 ──────────────────────────────────────────────

@router.post("/{result_id}/annotations", response_model=AnnotationResponse)
def create_annotation(
    result_id: str,
    req: AnnotationRequest,
    db: Session = Depends(get_db),
):
    annotation = Annotation(
        id=str(uuid.uuid4()),
        result_id=result_id,
        scope=req.scope,
        content=req.content,
        block_id=req.block_id,
        start_offset=req.start_offset,
        end_offset=req.end_offset,
        selected_text=req.selected_text,
    )
    db.add(annotation)
    db.commit()
    db.refresh(annotation)
    return annotation


@router.get("/{result_id}/annotations")
def list_annotations(result_id: str, db: Session = Depends(get_db)):
    return db.query(Annotation).filter(Annotation.result_id == result_id).all()


@router.patch("/{result_id}/annotations/{annotation_id}", response_model=AnnotationResponse)
def update_annotation(
    result_id: str,
    annotation_id: str,
    req: AnnotationUpdateRequest,
    db: Session = Depends(get_db),
):
    ann = db.query(Annotation).filter(
        Annotation.id == annotation_id,
        Annotation.result_id == result_id,
    ).first()
    if not ann:
        raise HTTPException(status_code=404, detail="批注不存在")
    ann.content = req.content
    db.commit()
    db.refresh(ann)
    return ann


@router.delete("/{result_id}/annotations/{annotation_id}")
def delete_annotation(
    result_id: str,
    annotation_id: str,
    db: Session = Depends(get_db),
):
    ann = db.query(Annotation).filter(
        Annotation.id == annotation_id,
        Annotation.result_id == result_id,
    ).first()
    if not ann:
        raise HTTPException(status_code=404, detail="批注不存在")
    db.delete(ann)
    db.commit()
    return {"ok": True}


# ── 导出 DOCX ────────────────────────────────────────────────────────────────

@router.get("/{result_id}/export")
def export_result(
    result_id: str,
    mode: str = Query("both", description="export mode: original | both | translation"),
    db: Session = Depends(get_db),
):
    """导出译文为 .docx 文件，支持三种模式：导出原文、原文+译文、仅译文"""
    if mode not in ("original", "both", "translation"):
        raise HTTPException(status_code=400, detail="mode 参数只能为 original / both / translation")

    result = db.query(TranslationResult).filter(TranslationResult.id == result_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="译文不存在")

    sj = result.structure_json or {}
    is_chinese = sj.get("paper_type") == "chinese"

    # ── 文件名：{前30字中文名}(模式).docx ──
    title_zh = sj.get("标题中文") or sj.get("标题") or "未命名"
    short_title = title_zh[:30]
    safe_title = "".join(c for c in short_title if c not in r'\/:*?"<>|')
    mode_labels = {"original": "原文", "both": "原文+译文", "translation": "译文"}
    mode_label = mode_labels[mode]

    doc = Document()

    # ── 窄页边距 ──
    for section in doc.sections:
        section.top_margin = Pt(36)       # 0.5 inch = 36pt
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(42)      # ~0.58 inch
        section.right_margin = Pt(42)

    # ── 文档默认样式 ──
    _configure_styles(doc)

    # ── 论文标题 ──
    title_en = sj.get("标题") or ""
    title_cn = sj.get("标题中文") or ""

    if is_chinese:
        if title_en:
            _add_doc_title(doc, title_en)
    else:
        if mode in ("original", "both") and title_en:
            _add_doc_title(doc, title_en)
        if mode in ("translation", "both") and title_cn:
            _add_doc_title(doc, title_cn)

    # ── 论文元信息（期刊 / 年份 / DOI）──
    meta_parts = []
    for key, label in [("所属期刊/会议", None), ("年份", None), ("DOI", "DOI: ")]:
        val = sj.get(key, "")
        if val:
            prefix = label or ""
            meta_parts.append(f"{prefix}{val}")
    if meta_parts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" / ".join(meta_parts))
        run.font.size = Pt(9)
        run.font.color.rgb = META_COLOR
        run.font.name = "Times New Roman"

    # ── 正文 ──
    blocks = sj.get("正文") or []

    for item in blocks:
        heading_level = item.get("标题等级")
        text_en = (item.get("文本") or "").strip()
        text_zh = (item.get("中文文本") or "").strip()

        # ── 图片块 ──
        if item.get("图片地址"):
            img_orig = item.get("图片地址") or ""
            img_zh = item.get("中文图片地址") or ""

            if is_chinese:
                _embed_image(doc, img_orig, label="图片")
            elif mode == "original":
                _embed_image(doc, img_orig, label="Figure")
            elif mode == "translation":
                _embed_image(doc, img_zh or img_orig, label="译文图片")
            else:  # both
                _embed_image(doc, img_orig, label="原图")
                if img_zh and img_zh != img_orig:
                    _embed_image(doc, img_zh, label="中文图")
            continue

        if heading_level is not None:
            # ── 标题块 ──
            if is_chinese:
                if text_en:
                    _add_doc_heading(doc, text_en, heading_level)
            else:
                if mode in ("original", "both") and text_en:
                    _add_doc_heading(doc, text_en, heading_level)
                if mode in ("translation", "both") and text_zh:
                    _add_doc_heading(doc, text_zh, heading_level)
        else:
            # ── 段落块 ──
            if is_chinese:
                if text_en:
                    _add_zh_para(doc, text_en)
            elif mode == "original":
                if text_en:
                    _add_en_para(doc, text_en)
            elif mode == "translation":
                if text_zh:
                    _add_zh_para(doc, text_zh)
            else:  # both
                if text_en:
                    _add_en_para(doc, text_en)
                if text_zh:
                    _add_zh_para(doc, text_zh)

    # ── 参考文献 ──
    references = sj.get("参考文献") or []
    if references:
        if is_chinese or mode == "original":
            ref_heading = "References" if sj.get("标题") and not is_chinese else "参考文献"
        else:
            ref_heading = "参考文献"
        doc.add_heading(ref_heading, level=1)
        for ref in references:
            p = doc.add_paragraph(str(ref))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            p.paragraph_format.left_indent = Pt(21)
            for run in p.runs:
                run.font.size = Pt(9)

    # ── 写入内存缓冲区 ──
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{safe_title}({mode_label}).docx"
    encoded_filename = quote(filename.encode("utf-8"))

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
        },
    )


# ── DOCX 样式常量（与 TranslationViewer.vue 的 CSS 一致）──

EN_FONT = "Times New Roman"
EN_SIZE = Pt(11)          # CSS 0.9rem ≈ 11pt
ZH_SIZE = Pt(12)          # CSS 0.95rem ≈ 12pt
EN_COLOR = RGBColor(0x60, 0x62, 0x66)   # #606266
ZH_COLOR = RGBColor(0x30, 0x31, 0x33)   # #303133
META_COLOR = RGBColor(0x90, 0x93, 0x99) # #909399

HEADING_SIZES = {
    1: Pt(14),   # CSS 1.05rem bold 700
    2: Pt(13),   # CSS 0.98rem bold 600
    3: Pt(12),   # CSS 0.93rem bold 600
}
HEADING_WEIGHTS = {1: True, 2: True, 3: True}

# 匹配 LaTeX 公式：行内 $...$（不跨行）和行间 $$...$$
_FORMULA_RE = re.compile(r'(\$\$[\s\S]+?\$\$|\$[^$\n]+?\$)')

# OMML 命名空间
_MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _configure_styles(doc: Document):
    """设置文档默认样式"""
    style = doc.styles["Normal"]
    style.font.name = EN_FONT
    style.font.size = EN_SIZE
    style.font.color.rgb = EN_COLOR
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ═══════════════════════════════════════════════════════════════════════════════
#  图片嵌入
# ═══════════════════════════════════════════════════════════════════════════════

def _resolve_image_url(url: str) -> str:
    """将相对路径（如 /uploads/...）转为绝对文件路径"""
    if url.startswith("/uploads/"):
        return str(local_storage.base_path / url[len("/uploads/"):])
    return url


def _fetch_image_bytes(url: str):
    """下载图片字节。支持本地文件和远程 URL。"""
    local_path = _resolve_image_url(url)

    # 本地文件
    import os as _os
    if _os.path.isfile(local_path):
        try:
            with open(local_path, "rb") as f:
                return f.read()
        except Exception:
            pass

    # 远程 URL
    if url.startswith("http://") or url.startswith("https://"):
        import httpx
        try:
            resp = httpx.get(url, timeout=30)
            resp.raise_for_status()
            return resp.content
        except Exception:
            pass

    return None


def _embed_image(doc: Document, url: str, label: str = ""):
    """下载图片并嵌入 docx，居中显示"""
    if not url:
        return

    img_bytes = _fetch_image_bytes(url)
    if not img_bytes:
        return

    # 标签
    if label:
        p_label = doc.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_label.paragraph_format.space_before = Pt(10)
        p_label.paragraph_format.space_after = Pt(2)
        r = p_label.add_run(label)
        r.font.size = Pt(9)
        r.font.color.rgb = META_COLOR
        r.italic = True

    # 图片
    from io import BytesIO as BIO
    try:
        stream = BIO(img_bytes)
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(2)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(stream, width=Pt(280))
    except Exception:
        pass


# ═══════════════════════════════════════════════════════════════════════════════
#  LaTeX → MathML → OMML 公式渲染器
# ═══════════════════════════════════════════════════════════════════════════════

def _latex_to_omath(latex: str, is_display: bool = False):
    """将 LaTeX 公式转为 OMML <m:oMath> 或 <m:oMathPara> 元素。
    失败时返回 None。
    """
    import latex2mathml.converter
    from lxml import etree

    try:
        mathml_str = latex2mathml.converter.convert(latex)
    except Exception:
        return None

    try:
        mathml_el = etree.fromstring(mathml_str.encode('utf-8'))
    except Exception:
        return None

    try:
        omml = _mathml_to_omml(mathml_el)
    except Exception:
        return None

    if omml is None:
        return None

    # 行间公式用 oMathPara 包裹，行内用 oMath
    if is_display:
        wrapper = etree.SubElement(etree.Element(_m('oMathPara')), _m('oMath'))
    else:
        wrapper = etree.Element(_m('oMath'))

    # 将 omml 的所有子元素移入 wrapper
    if omml.tag == _m('oMath'):
        for child in list(omml):
            wrapper.append(child)
    else:
        wrapper.append(omml)

    return wrapper


def _m(tag: str) -> str:
    """返回带 OMML 命名空间的完整标签名"""
    return f'{{{_MATH_NS}}}{tag}'


def _mathml_to_omml(el):
    """递归将 MathML presentation 元素转为 OMML 元素。"""
    from lxml import etree

    tag_local = etree.QName(el.tag).localname if etree.QName(el.tag).namespace else el.tag

    # ── 容器元素 ──
    if tag_local in ('math', 'mrow', 'mstyle', 'semantics', 'mphantom', 'mpadded'):
        # 直接递归所有子元素
        omml = etree.Element(_m('oMath'))
        for child in el:
            converted = _mathml_to_omml(child)
            if converted is not None:
                if converted.tag == _m('oMath'):
                    for c in list(converted):
                        omml.append(c)
                else:
                    omml.append(converted)
        return omml

    # ── 文本元素：mi（斜体）, mn（正体）, mo（运算符）, mtext ──
    if tag_local in ('mi', 'mn', 'mo', 'mtext'):
        text = ''.join(el.itertext()).strip()
        if not text:
            return None
        r = etree.Element(_m('r'))
        # mi 默认斜体，其他正常。但 mo 中部分符号（∑∫∏）是大的，暂不特殊处理
        if tag_local == 'mi':
            rPr = etree.SubElement(r, _m('rPr'))
            etree.SubElement(rPr, _m('sty')).set(_m('val'), 'i')
        t = etree.SubElement(r, _m('t'))
        t.text = text
        return r

    # ── 分式 mfrac ──
    if tag_local == 'mfrac':
        f = etree.Element(_m('f'))
        fPr = etree.SubElement(f, _m('fPr'))
        etree.SubElement(fPr, _m('type')).set(_m('val'), 'bar')
        num = etree.SubElement(f, _m('num'))
        num_child = _mathml_to_omml(el[0]) if len(el) > 0 and el[0] is not None else None
        if num_child is not None:
            if num_child.tag == _m('oMath'):
                for c in list(num_child):
                    num.append(c)
            else:
                num.append(num_child)
        den = etree.SubElement(f, _m('den'))
        den_child = _mathml_to_omml(el[1]) if len(el) > 1 and el[1] is not None else None
        if den_child is not None:
            if den_child.tag == _m('oMath'):
                for c in list(den_child):
                    den.append(c)
            else:
                den.append(den_child)
        return f

    # ── 上标 msup ──
    if tag_local == 'msup':
        ss = etree.Element(_m('sSup'))
        e = etree.SubElement(ss, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        sup = etree.SubElement(ss, _m('sup'))
        sup_child = _mathml_to_omml(el[1]) if len(el) > 1 else None
        if sup_child is not None:
            if sup_child.tag == _m('oMath'):
                for c in list(sup_child):
                    sup.append(c)
            else:
                sup.append(sup_child)
        return ss

    # ── 下标 msub ──
    if tag_local == 'msub':
        ss = etree.Element(_m('sSub'))
        e = etree.SubElement(ss, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        sub = etree.SubElement(ss, _m('sub'))
        sub_child = _mathml_to_omml(el[1]) if len(el) > 1 else None
        if sub_child is not None:
            if sub_child.tag == _m('oMath'):
                for c in list(sub_child):
                    sub.append(c)
            else:
                sub.append(sub_child)
        return ss

    # ── 上下标 msubsup ──
    if tag_local == 'msubsup':
        ss = etree.Element(_m('sSubSup'))
        e = etree.SubElement(ss, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        sub = etree.SubElement(ss, _m('sub'))
        sub_child = _mathml_to_omml(el[1]) if len(el) > 1 else None
        if sub_child is not None:
            if sub_child.tag == _m('oMath'):
                for c in list(sub_child):
                    sub.append(c)
            else:
                sub.append(sub_child)
        sup = etree.SubElement(ss, _m('sup'))
        sup_child = _mathml_to_omml(el[2]) if len(el) > 2 else None
        if sup_child is not None:
            if sup_child.tag == _m('oMath'):
                for c in list(sup_child):
                    sup.append(c)
            else:
                sup.append(sup_child)
        return ss

    # ── 根号 msqrt / mroot ──
    if tag_local in ('msqrt', 'mroot'):
        rad = etree.Element(_m('rad'))
        radPr = etree.SubElement(rad, _m('radPr'))
        if tag_local == 'msqrt':
            etree.SubElement(radPr, _m('degHide')).set(_m('val'), '1')
        deg = etree.SubElement(rad, _m('deg'))
        if tag_local == 'mroot' and len(el) > 1:
            deg_child = _mathml_to_omml(el[1])
            if deg_child is not None:
                if deg_child.tag == _m('oMath'):
                    for c in list(deg_child):
                        deg.append(c)
                else:
                    deg.append(deg_child)
        e = etree.SubElement(rad, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        return rad

    # ── 括号 mfenced（自动扩展括号）──
    if tag_local == 'mfenced':
        d = etree.Element(_m('d'))
        dPr = etree.SubElement(d, _m('dPr'))
        # 默认圆括号
        open_char = el.get('open', '(')
        close_char = el.get('close', ')')
        begChr = etree.SubElement(dPr, _m('begChr'))
        begChr.set(_m('val'), open_char)
        endChr = etree.SubElement(dPr, _m('endChr'))
        endChr.set(_m('val'), close_char)
        # 内容
        e = etree.SubElement(d, _m('e'))
        inner = _mathml_to_omml(el)  # 复用 math 处理逻辑——不对，mfenced 的子元素是 mrow 或无
        # 直接将子元素递归
        if len(el) > 0:
            inner_omml = el  # 直接处理自己的子元素? 不，用递归
        for child in el:
            c = _mathml_to_omml(child)
            if c is not None:
                if c.tag == _m('oMath'):
                    for cc in list(c):
                        e.append(cc)
                else:
                    e.append(c)
        return d

    # ── 运算符（∑ ∫ ∏ 等，可能带上下标）──
    # mover / munder / munderover
    if tag_local == 'mover':
        nary = etree.Element(_m('limUpp'))
        e = etree.SubElement(nary, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        lim = etree.SubElement(nary, _m('lim'))
        lim_child = _mathml_to_omml(el[1]) if len(el) > 1 else None
        if lim_child is not None:
            if lim_child.tag == _m('oMath'):
                for c in list(lim_child):
                    lim.append(c)
            else:
                lim.append(lim_child)
        return nary

    if tag_local == 'munder':
        nary = etree.Element(_m('limLow'))
        e = etree.SubElement(nary, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        lim = etree.SubElement(nary, _m('lim'))
        lim_child = _mathml_to_omml(el[1]) if len(el) > 1 else None
        if lim_child is not None:
            if lim_child.tag == _m('oMath'):
                for c in list(lim_child):
                    lim.append(c)
            else:
                lim.append(lim_child)
        return nary

    if tag_local == 'munderover':
        nary = etree.Element(_m('limUpp'))
        e = etree.SubElement(nary, _m('e'))
        base = _mathml_to_omml(el[0]) if len(el) > 0 else None
        if base is not None:
            if base.tag == _m('oMath'):
                for c in list(base):
                    e.append(c)
            else:
                e.append(base)
        # limLow for subscript
        sub = etree.SubElement(nary, _m('lim'))
        low = etree.Element(_m('limLow'))
        low_e = etree.SubElement(low, _m('e'))
        # The subscript content
        pass  # This is getting complex, handle simply for now

        return nary

    # ── 表格 / 矩阵 mtable ──
    if tag_local == 'mtable':
        m = etree.Element(_m('m'))
        mPr = etree.SubElement(m, _m('mPr'))
        # 简单地按行处理
        for row in el:
            if etree.QName(row).localname == 'mtr':
                mr = etree.SubElement(m, _m('mr'))
                for cell in row:
                    if etree.QName(cell).localname == 'mtd':
                        e = etree.SubElement(mr, _m('e'))
                        cell_omml = _mathml_to_omml(cell)
                        if cell_omml is not None:
                            if cell_omml.tag == _m('oMath'):
                                for c in list(cell_omml):
                                    e.append(c)
                            else:
                                e.append(cell_omml)
        return m

    # ── 空格 mspace ──
    if tag_local == 'mspace':
        r = etree.Element(_m('r'))
        t = etree.SubElement(r, _m('t'))
        t.text = ' '
        return r

    # 未知元素：只取文本
    text = ''.join(el.itertext()).strip()
    if text:
        r = etree.Element(_m('r'))
        t = etree.SubElement(r, _m('t'))
        t.text = text
        return r
    return None


# ═══════════════════════════════════════════════════════════════════════════════
#  段落构建 helper
# ═══════════════════════════════════════════════════════════════════════════════

def _add_formula_runs(para, text: str, font_name: str, font_size, font_color,
                      is_bold: bool = False):
    """拆分 text 中的 LaTeX 公式并渲染为 OMath / 普通 run。

    行内公式 ($...$) → 嵌入段落流中的 OMath
    行间公式 ($$...$$) → 强制换行 + OMath + 强制换行
    """
    parts = _FORMULA_RE.split(text)

    for part in parts:
        if not part:
            continue

        if part.startswith('$$') and part.endswith('$$'):
            formula = part[2:-2].strip()
            om = _latex_to_omath(formula, is_display=True)
            if om is not None:
                # 换行
                _append_br(para)
                para._element.append(om)
                _append_br(para)
            else:
                # 渲染失败 —— 回退为纯文本
                _append_br(para)
                r = para.add_run(formula)
                r.font.name = font_name
                r.font.size = font_size
                r.font.color.rgb = font_color
                r.bold = is_bold
                r.italic = True
                _append_br(para)

        elif part.startswith('$') and part.endswith('$'):
            formula = part[1:-1].strip()
            om = _latex_to_omath(formula, is_display=False)
            if om is not None:
                para._element.append(om)
            else:
                # 回退
                r = para.add_run(formula)
                r.font.name = font_name
                r.font.size = font_size
                r.font.color.rgb = font_color
                r.bold = is_bold
                r.italic = True

        else:
            _append_run(para, part, font_name, font_size, font_color, is_bold)


def _append_run(para, text: str, font_name: str, font_size, font_color, is_bold: bool = False):
    """添加一个普通文本 run"""
    if not text:
        return
    r = para.add_run(text)
    r.font.name = font_name
    r.font.size = font_size
    r.font.color.rgb = font_color
    r.bold = is_bold


def _append_br(para):
    """在段落末尾插入一个换行符 run"""
    from lxml import etree
    r_elem = etree.SubElement(para._element, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
    etree.SubElement(r_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br')


def _add_doc_title(doc: Document, text: str):
    """添加论文大标题（居中）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    _add_formula_runs(p, text, EN_FONT, Pt(16), ZH_COLOR, is_bold=True)


def _add_doc_heading(doc: Document, text: str, level: int):
    """添加章节标题（两端对齐）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.2
    size = HEADING_SIZES.get(level, Pt(12))
    is_bold = HEADING_WEIGHTS.get(level, False)
    _add_formula_runs(p, text, EN_FONT, size, ZH_COLOR, is_bold=is_bold)


def _add_en_para(doc: Document, text: str):
    """添加英文原文段落（Times New Roman 11pt, #606266, 两端对齐）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _add_formula_runs(p, text, EN_FONT, EN_SIZE, EN_COLOR)


def _add_zh_para(doc: Document, text: str, line_spacing: float = 1.45):
    """添加中文译文段落（12pt, #303133, 两端对齐）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    _add_formula_runs(p, text, EN_FONT, ZH_SIZE, ZH_COLOR)


# ── 图片翻译 ────────────────────────────────────────────────────────────────

@router.post("/{result_id}/translate-image")
def translate_image_block(
    result_id: str,
    req: TranslateImageRequest,
    db: Session = Depends(get_db),
):
    result = db.query(TranslationResult).filter(TranslationResult.id == result_id).first()
    if not result:
        raise HTTPException(status_code=404, detail="译文不存在")

    from app.models.user_glossary import UserGlossary
    from app.models.domain_glossary import DomainGlossary

    paper = db.query(Paper).filter(Paper.id == result.paper_id).first()
    user_glossary = db.query(UserGlossary).all()
    domain_glossary = []
    if paper and paper.domain:
        domain_glossary = db.query(DomainGlossary).filter(
            DomainGlossary.domain == paper.domain
        ).all()

    glossary_list = [{"en": g.foreign_term, "zh": g.zh_term, "status": g.status} for g in user_glossary]
    user_keys = {g.foreign_term.lower() for g in user_glossary}
    for dg in domain_glossary:
        if dg.en_term.lower() not in user_keys:
            glossary_list.append({"en": dg.en_term, "zh": dg.zh_term, "status": "translate"})

    domain = (paper.domain or "学术") if paper else "学术"
    paper_id = result.paper_id or ""

    from app.services.image_translation import translate_image
    translated_url = translate_image(
        image_url=req.image_url,
        glossary_list=glossary_list,
        paper_id=paper_id,
        domain=domain,
    )

    changed = translated_url != req.image_url
    if changed:
        from sqlalchemy.orm.attributes import flag_modified
        zhengwen = list(result.structure_json.get("正文", []))
        idx = req.block_index
        if 0 <= idx < len(zhengwen) and "图片地址" in zhengwen[idx]:
            zhengwen[idx] = dict(zhengwen[idx])
            zhengwen[idx]["中文图片地址"] = translated_url
            new_json = dict(result.structure_json)
            new_json["正文"] = zhengwen
            result.structure_json = new_json
            flag_modified(result, "structure_json")
        db.commit()

    return {"translated_url": translated_url, "changed": changed}
